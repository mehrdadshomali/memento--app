#!/usr/bin/env python3
"""
Memento - Professional A3 Poster Generator (DOCX)
Ultra-premium academic poster with warm brown/amber/cream palette
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml
import os

# ===== COLOR PALETTE =====
DARK_BROWN = "3B2F2A"
AMBER = "B07D4F"
AMBER_LIGHT = "E9C9A0"
CREAM = "FAF7F3"
CREAM_DARK = "F0E6DA"
TEXT_DARK = "2C2420"
TEXT_MID = "4A3F38"
TEXT_LIGHT = "8A7A6E"
WHITE = "FFFFFF"
CHECK_GREEN = "7B9E6F"
ARCH_GREEN = "E8EFE4"
ARCH_BROWN = "E4DDD5"
BORDER_CREAM = "E0D5C8"

def set_cell_shading(cell, color):
    """Apply background color to a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    tcPr.append(shading)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set individual cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            color, sz = val
            border = parse_xml(
                f'<w:{side} {nsdecls("w")} w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            )
            tcBorders.append(border)
    tcPr.append(tcBorders)

def set_cell_margins(cell, top=0, bottom=0, left=100, right=100):
    """Set cell padding/margins in TWIPs."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    margins = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:start w:w="{left}" w:type="dxa"/>'
        f'<w:end w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(margins)

def set_table_borders(table, color="FFFFFF", sz="0"):
    """Remove or set all table borders."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="0" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def set_col_width(col, width_cm):
    """Set column width."""
    for cell in col.cells:
        cell.width = Cm(width_cm)

def make_run(para, text, size=10, color=TEXT_MID, bold=False, italic=False, font_name="Calibri"):
    """Create a styled run."""
    r = para.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(int(color[:2],16), int(color[2:4],16), int(color[4:],16))
    r.bold = bold
    r.italic = italic
    r.font.name = font_name
    return r

def set_para_spacing(para, before=0, after=0, line_spacing=1.15):
    """Set paragraph spacing."""
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing

def add_section_header(cell, num, tag, title, dark=False):
    """Add a beautifully formatted section header with number and tag."""
    # Number + Tag line
    p1 = cell.add_paragraph()
    set_para_spacing(p1, before=2, after=1)
    make_run(p1, num, size=16, color=AMBER if not dark else AMBER_LIGHT, bold=True)
    make_run(p1, f"   {tag}", size=7, color=TEXT_LIGHT if not dark else "AAAAAA", bold=True)
    
    # Title
    p2 = cell.add_paragraph()
    set_para_spacing(p2, before=0, after=4)
    make_run(p2, title, size=13, color=TEXT_DARK if not dark else WHITE, bold=True)
    
    # Divider line
    p3 = cell.add_paragraph()
    set_para_spacing(p3, before=0, after=4)
    make_run(p3, "━" * 35, size=5, color=AMBER if not dark else "5A4A3E")

def add_body_text(cell, text, dark=False, size=9):
    """Add body paragraph text."""
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_spacing(p, before=0, after=3, line_spacing=1.3)
    make_run(p, text, size=size, color=TEXT_MID if not dark else "CCCCCC")

def add_check_item(cell, text):
    """Add a checkmark list item."""
    p = cell.add_paragraph()
    set_para_spacing(p, before=0, after=2, line_spacing=1.2)
    make_run(p, "  ✓  ", size=10, color=CHECK_GREEN, bold=True)
    make_run(p, text, size=9, color=TEXT_MID)

def add_screenshot_box(cell, icon, label, filename):
    """Add a screenshot placeholder box using a nested table."""
    # Create a mini table for the placeholder
    p = cell.add_paragraph()
    set_para_spacing(p, before=2, after=2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    make_run(p, f"┌{'─' * 28}┐", size=8, color=BORDER_CREAM)
    
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p2, before=0, after=0)
    make_run(p2, icon, size=16, color=TEXT_LIGHT)
    
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p3, before=1, after=0)
    make_run(p3, label, size=8, color=TEXT_LIGHT, bold=True)
    
    p4 = cell.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p4, before=0, after=1)
    make_run(p4, filename, size=7, color=BORDER_CREAM, italic=True)
    
    p5 = cell.add_paragraph()
    set_para_spacing(p5, before=0, after=2)
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    make_run(p5, f"└{'─' * 28}┘", size=8, color=BORDER_CREAM)

def add_tech_pill(para, text):
    """Add a tech badge/pill."""
    make_run(para, f"  {text}  ", size=8, color=TEXT_MID, bold=True)
    make_run(para, "  •  ", size=7, color=BORDER_CREAM)

def add_future_tag(para, text, is_last=False):
    """Add a future work tag."""
    make_run(para, f" [ {text} ] ", size=8, color="AAAAAA")
    if not is_last:
        make_run(para, "  ", size=7, color="555555")

# ============================================================
# BUILD DOCUMENT
# ============================================================
doc = Document()

# Clear default styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.font.color.rgb = RGBColor(0x2c, 0x24, 0x20)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)

# A3 Portrait
section = doc.sections[0]
section.page_height = Cm(42.0)
section.page_width = Cm(29.7)
section.orientation = WD_ORIENT.PORTRAIT
section.left_margin = Cm(1.2)
section.right_margin = Cm(1.2)
section.top_margin = Cm(1.2)
section.bottom_margin = Cm(1.2)

# ============================================================
# HEADER BANNER
# ============================================================
header_tbl = doc.add_table(rows=1, cols=1)
header_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(header_tbl)
hc = header_tbl.cell(0, 0)
set_cell_shading(hc, DARK_BROWN)
set_cell_margins(hc, top=200, bottom=200, left=300, right=300)
hc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# University line
p = hc.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
make_run(p, "🎓  ", size=12, color=WHITE)
make_run(p, "TEKİRDAĞ NAMIK KEMAL ÜNİVERSİTESİ", size=10, color=WHITE, bold=True)
make_run(p, "  |  Çorlu Mühendislik Fakültesi · Bilgisayar Mühendisliği", size=8, color="999999")

# mHEALTH tags (right side effect via tabs)
p_tags = hc.add_paragraph()
p_tags.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_para_spacing(p_tags, before=0, after=6)
make_run(p_tags, "  mHEALTH  ", size=7, color="999999", bold=True)
make_run(p_tags, "   ", size=7, color="555555")
make_run(p_tags, "  PROJE-II · 2026  ", size=7, color="999999", bold=True)

# Title
p_title = hc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_para_spacing(p_title, before=4, after=2)
r_title = make_run(p_title, "Memento.", size=42, color=WHITE, bold=True, italic=True, font_name="Georgia")

# Subtitle
p_sub = hc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_para_spacing(p_sub, before=0, after=8)
make_run(p_sub, "Alzheimer Hastaları İçin Terapötik Hafıza Egzersiz ve\nGüvenlik Odaklı Mobil Uygulama", size=13, color="CCCCCC")

# Authors
p_auth = hc.add_paragraph()
p_auth.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_para_spacing(p_auth, before=4, after=2)
make_run(p_auth, "Hazırlayanlar\n", size=7, color="777777", bold=True)
make_run(p_auth, "Mehrdad Shomali", size=9, color=AMBER_LIGHT, bold=True)
make_run(p_auth, " · 2210656637\n", size=8, color="AAAAAA")
make_run(p_auth, "Kerim Can Güzel", size=9, color=AMBER_LIGHT, bold=True)
make_run(p_auth, " · 2220656047\n", size=8, color="AAAAAA")
make_run(p_auth, "Efekan Çetin", size=9, color=AMBER_LIGHT, bold=True)
make_run(p_auth, " · 2220656034", size=8, color="AAAAAA")

# Advisor
p_adv = hc.add_paragraph()
p_adv.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_para_spacing(p_adv, before=4, after=0)
make_run(p_adv, "Danışman\n", size=7, color="777777", bold=True)
make_run(p_adv, "Dr. Öğr. Üyesi Erkan Özhan", size=10, color=AMBER_LIGHT, bold=True)

# ============================================================
# STATS BAR
# ============================================================
doc.add_paragraph()  # tiny spacer
stats_tbl = doc.add_table(rows=1, cols=4)
stats_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(stats_tbl)

stats_data = [
    ("55M+", "kişi demans ile yaşıyor (WHO)"),
    ("%60–70", "vakanın nedeni Alzheimer"),
    ("100%", "çevrimdışı çalışabilen mimari"),
    ("6", "bütünleşik modül · iOS & Android"),
]
for i, (val, label) in enumerate(stats_data):
    c = stats_tbl.cell(0, i)
    set_cell_shading(c, WHITE)
    set_cell_margins(c, top=120, bottom=120, left=150, right=150)
    set_cell_borders(c, 
        top=(BORDER_CREAM, "4"), 
        bottom=(BORDER_CREAM, "4"), 
        left=(BORDER_CREAM, "4"), 
        right=(BORDER_CREAM, "4"))
    
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, before=2, after=1)
    make_run(p, val, size=22, color=DARK_BROWN, bold=True)
    
    p2 = c.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p2, before=0, after=2)
    make_run(p2, label, size=7, color=TEXT_LIGHT)

# ============================================================
# MAIN 3-COLUMN GRID
# ============================================================
doc.add_paragraph()  # spacer

main_tbl = doc.add_table(rows=1, cols=3)
main_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(main_tbl)

# Set column widths
for cell in main_tbl.columns[0].cells: cell.width = Cm(8.8)
for cell in main_tbl.columns[1].cells: cell.width = Cm(9.2)
for cell in main_tbl.columns[2].cells: cell.width = Cm(8.8)

left = main_tbl.cell(0, 0)
mid = main_tbl.cell(0, 1)
right = main_tbl.cell(0, 2)

# Apply cream background and margins
for c in [left, mid, right]:
    set_cell_shading(c, CREAM)
    set_cell_margins(c, top=150, bottom=150, left=200, right=200)
    set_cell_borders(c,
        top=(BORDER_CREAM, "6"),
        bottom=(BORDER_CREAM, "6"),
        left=(BORDER_CREAM, "2"),
        right=(BORDER_CREAM, "2"))
    c.vertical_alignment = WD_ALIGN_VERTICAL.TOP

# ============================================================
# LEFT COLUMN
# ============================================================
left.paragraphs[0].text = ""  # clear default

# 01 - Problem & Motivasyon
add_section_header(left, "01", "PROBLEM", "Problem & Motivasyon")
add_body_text(left, "Alzheimer; hafızayı, dili ve yön bulmayı geri döndürülemez biçimde zayıflatan ilerleyici bir hastalıktır. Hastalar yakınlarını tanıyamaz, ilaç rutinlerini unutur ve evden uzaklaşıp kaybolma riski taşır.")
add_body_text(left, "Mevcut uygulamalar yaşlı kullanıcıya uygun olmayan karmaşık arayüzlere sahip olup çoğunlukla tek bir özelliğe odaklıdır.")

# Spacer
p_sp = left.add_paragraph()
set_para_spacing(p_sp, before=6, after=0)
make_run(p_sp, "", size=2, color=CREAM)

# 02 - Amaç
add_section_header(left, "02", "HEDEFLER", "Amaç")
add_check_item(left, "Kişiselleştirilebilir bilişsel hafıza egzersizleri")
add_check_item(left, "Coğrafi çit ile evden uzaklaşmayı tespit ve yönlendirme")
add_check_item(left, "İlaç ve rutinler için zamanlanmış hatırlatmalar")
add_check_item(left, "Bakıcıya nesnel ilerleme analitiği")
add_check_item(left, "WCAG uyumlu, yaşlı dostu erişilebilir arayüz")

# Spacer
p_sp = left.add_paragraph()
set_para_spacing(p_sp, before=6, after=0)
make_run(p_sp, "", size=2, color=CREAM)

# 03 - Sistem Mimarisi
add_section_header(left, "03", "MİMARİ", "Sistem Mimarisi")

# Architecture layers
layers = [
    ("Sunum Katmanı", "React Native ekranları & UI bileşenleri", CREAM_DARK),
    ("İş Mantığı", "Context API · useReducer · Custom Hooks", ARCH_GREEN),
    ("Veri / Servis", "Supabase (PostgreSQL) · AsyncStorage önbellek", ARCH_BROWN),
]
for name, desc, color in layers:
    p_l = left.add_paragraph()
    p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p_l, before=2, after=0)
    make_run(p_l, f"┌{'─' * 38}┐", size=7, color=BORDER_CREAM)
    
    p_n = left.add_paragraph()
    p_n.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p_n, before=0, after=0)
    make_run(p_n, name, size=9, color=TEXT_DARK, bold=True, italic=True)
    
    p_d = left.add_paragraph()
    p_d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p_d, before=0, after=0)
    make_run(p_d, desc, size=7, color=TEXT_LIGHT)
    
    p_b = left.add_paragraph()
    p_b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p_b, before=0, after=1)
    make_run(p_b, f"└{'─' * 38}┘", size=7, color=BORDER_CREAM)
    
    # Arrow
    p_a = left.add_paragraph()
    p_a.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p_a, before=0, after=0)
    make_run(p_a, "↓", size=10, color=BORDER_CREAM)

# Tech pills
p_tech = left.add_paragraph()
set_para_spacing(p_tech, before=6, after=2)
for t in ["React Native", "Expo SDK 54", "TypeScript", "Supabase", "expo-location"]:
    add_tech_pill(p_tech, t)

# Spacer
p_sp = left.add_paragraph()
set_para_spacing(p_sp, before=6, after=0)
make_run(p_sp, "", size=2, color=CREAM)

# 04 - Erişilebilirlik
add_section_header(left, "04", "ALTYAPI", "Erişilebilirlik & Çoklu Dil")
add_body_text(left, "Beynin aşırı uyarılmasını önlemek için pastel tonlar (Sage, Wood, Blue) tercih edilmiştir. İnce motor beceri kaybına karşı buton yükseklikleri minimum 56px yapılarak devasa dokunma hedefleri sağlanmıştır.")
add_body_text(left, "Proje baştan itibaren uluslararasılaştırma (i18n) destekli kodlanmıştır. Ayarlar üzerinden anında Türkçe/İngilizce dil değişikliği yapılabilir.")


# ============================================================
# MIDDLE COLUMN
# ============================================================
mid.paragraphs[0].text = ""

# 05 - Çekirdek Algoritmalar
add_section_header(mid, "05", "YÖNTEM", "Çekirdek Algoritmalar")

# Algorithm boxes side by side (simulated)
p_algo = mid.add_paragraph()
p_algo.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p_algo, before=2, after=0)
make_run(p_algo, "🔀  Fisher-Yates", size=11, color=TEXT_DARK, bold=True)
make_run(p_algo, "          ", size=10, color=CREAM)
make_run(p_algo, "📍  Haversine", size=11, color=TEXT_DARK, bold=True)

p_algo2 = mid.add_paragraph()
p_algo2.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p_algo2, before=1, after=1)
make_run(p_algo2, "O(N)", size=14, color=AMBER, bold=True)
make_run(p_algo2, "                        ", size=10, color=CREAM)
make_run(p_algo2, "≤100 m", size=14, color=AMBER, bold=True)

p_algo3 = mid.add_paragraph()
p_algo3.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p_algo3, before=0, after=4)
make_run(p_algo3, "Kartları eşit olasılıklı karıştırır", size=7, color=TEXT_LIGHT)
make_run(p_algo3, "       ", size=7, color=CREAM)
make_run(p_algo3, "Jeodezik mesafe, güvenli eşik", size=7, color=TEXT_LIGHT)

add_body_text(mid, "Arka plan konum görevi cihaz kapalıyken bile çalışır; pil için zaman (15 dk) ve mesafe (50 m) filtreleriyle optimize edilmiştir.", size=8)

# Spacer
p_sp = mid.add_paragraph()
set_para_spacing(p_sp, before=6, after=0)
make_run(p_sp, "", size=2, color=CREAM)

# 06 - Bilişsel & İşitsel Rehabilitasyon
add_section_header(mid, "06", "MODÜLLER", "Bilişsel & İşitsel Rehabilitasyon")

# Module: Aile Albümü
add_screenshot_box(mid, "🖼️", "AİLE ALBÜMÜ", "FamilyAlbumScreen.tsx")
p_m1 = mid.add_paragraph()
set_para_spacing(p_m1, before=0, after=2)
make_run(p_m1, "🖼️  Aile Albümü & Hafıza Egzersizi", size=9, color=TEXT_DARK, bold=True)
add_body_text(mid, "Yakınların fotoğrafları soru-cevaba dönüşür; doğru cevap kutlanır, yanlışta nazik ipucu verilir.", size=8)

# Module: Sesli Mesajlar
add_screenshot_box(mid, "🎙️", "SESLİ MESAJLAR", "VoiceMessagesScreen.tsx")
p_m2 = mid.add_paragraph()
set_para_spacing(p_m2, before=0, after=2)
make_run(p_m2, "🎙️  Sesli Mesajlar", size=9, color=TEXT_DARK, bold=True)
add_body_text(mid, "Sevdiklerinin kendi sesinden kayıtlar; işitsel hafıza daha uzun korunur.", size=8)

# Module: Coğrafi Çit
add_screenshot_box(mid, "🗺️", "GÜVENLİK HARİTASI", "SafetyScreen.tsx")
p_m3 = mid.add_paragraph()
set_para_spacing(p_m3, before=0, after=2)
make_run(p_m3, "🗺️  Coğrafi Çit & Güvenlik", size=9, color=TEXT_DARK, bold=True)
add_body_text(mid, "Ev konumu kaydedilir; güvenli bölge dışına çıkınca \"Eve Yol Tarifi\" otomatik açılır.", size=8)

# Spacer
p_sp = mid.add_paragraph()
set_para_spacing(p_sp, before=4, after=0)
make_run(p_sp, "", size=2, color=CREAM)

# 07 - Onboarding
add_section_header(mid, "07", "KULLANICI DENEYİMİ", "Onboarding & Karşılama")
add_screenshot_box(mid, "🚀", "SPLASH SCREEN", "AnimatedSplashScreen.tsx")
add_body_text(mid, "Hasta kafa karışıklığı yaşamasın diye özel, yavaş ve yatıştırıcı açılış ekranı animasyonu tasarlanmıştır.", size=8)


# ============================================================
# RIGHT COLUMN
# ============================================================
right.paragraphs[0].text = ""

# 08 - Rutin Takibi & Bakıcı
add_section_header(right, "08", "MODÜLLER", "Rutin Takibi & Bakıcı Paneli")

# Rutin
add_screenshot_box(right, "⏰", "RUTİNLER", "RoutineScreen.tsx")
p_r1 = right.add_paragraph()
set_para_spacing(p_r1, before=0, after=2)
make_run(p_r1, "💊  Günlük Rutin & İlaç", size=9, color=TEXT_DARK, bold=True)
add_body_text(right, "Yedi kategoride zamanlanmış bildirimler hastayı saatinde uyarır.", size=8)

# Analitik
add_screenshot_box(right, "📊", "ANALİTİK", "AnalyticsScreen.tsx")
p_r2 = right.add_paragraph()
set_para_spacing(p_r2, before=0, after=2)
make_run(p_r2, "📊  Bakıcı Analitiği", size=9, color=TEXT_DARK, bold=True)
add_body_text(right, "İlaç uyumu ve hafıza testi eğilimleri grafiklerle; doktor randevuları için değerli veri.", size=8)

# Offline-first note
p_off = right.add_paragraph()
set_para_spacing(p_off, before=4, after=4)
make_run(p_off, "  ⚡ ", size=9, color=AMBER)
make_run(p_off, "Çevrimdışı öncelikli: ", size=8, color=TEXT_DARK, bold=True)
make_run(p_off, "veriler Supabase ile senkronlanır, yerel önbelleğe aynılanır; internet kesilse de uygulama çalışır.", size=8, color=TEXT_MID)

# Spacer
p_sp = right.add_paragraph()
set_para_spacing(p_sp, before=4, after=0)
make_run(p_sp, "", size=2, color=CREAM)

# 09 - Supabase DB
add_section_header(right, "09", "VERİTABANI", "Supabase (PostgreSQL) Şeması")
add_body_text(right, "Firebase'in doküman yapısı yerine, ilişkisel bütünlük için Supabase tercih edilmiştir. Medya dosyaları Storage Bucket'a yüklenir, DB'de yalnızca genel URL tutulur.", size=8)

# DB Table
db_data = [
    ("Tablo Adı", "FK Bağıntısı"),
    ("profiles", "— (Ana Tablo)"),
    ("memory_cards", "profile_id → profiles(id)"),
    ("routines", "profile_id → profiles(id)"),
    ("routine_completions", "routine_id → routines(id)"),
    ("game_sessions", "profile_id → profiles(id)"),
    ("safety_profiles", "profile_id → profiles(id)"),
]
for i, (tbl_name, fk) in enumerate(db_data):
    p_db = right.add_paragraph()
    set_para_spacing(p_db, before=0, after=1)
    if i == 0:
        # Header row
        make_run(p_db, f"  {tbl_name}", size=8, color=TEXT_DARK, bold=True)
        make_run(p_db, f"    {fk}", size=8, color=TEXT_DARK, bold=True)
        p_div = right.add_paragraph()
        set_para_spacing(p_div, before=0, after=1)
        make_run(p_div, "━" * 40, size=5, color=AMBER)
    else:
        make_run(p_db, f"  {tbl_name}", size=8, color=AMBER, bold=True, font_name="Courier New")
        make_run(p_db, f"    {fk}", size=7, color=TEXT_LIGHT, font_name="Courier New")

# Spacer
p_sp = right.add_paragraph()
set_para_spacing(p_sp, before=6, after=0)
make_run(p_sp, "", size=2, color=CREAM)

# ============================================================
# 10 - SONUÇ (DARK CARD - simulated in the same cell)
# ============================================================
# Separator
p_sep = right.add_paragraph()
set_para_spacing(p_sep, before=4, after=4)
make_run(p_sep, "━" * 40, size=5, color=DARK_BROWN)

add_section_header(right, "10", "SONUÇ", "Sonuç & Gelecek Çalışmalar", dark=False)
add_body_text(right, "Memento; bilişsel rehabilitasyon, güvenlik ve bakıcı desteğini tek bütünleşik uygulamada birleştirir ve belirlenen tüm işlevsel hedeflere ulaşır.")

# Future tags
p_ft = right.add_paragraph()
set_para_spacing(p_ft, before=4, after=2)
add_future_tag(p_ft, "SOS acil çağrı")
add_future_tag(p_ft, "Yapay zekâ yüz tanıma")
add_future_tag(p_ft, "Akıllı saat · düşme algılama", is_last=True)


# ============================================================
# FOOTER BANNER
# ============================================================
doc.add_paragraph()  # spacer

footer_tbl = doc.add_table(rows=1, cols=1)
footer_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(footer_tbl)
fc = footer_tbl.cell(0, 0)
set_cell_shading(fc, DARK_BROWN)
set_cell_margins(fc, top=100, bottom=100, left=300, right=300)

pf = fc.paragraphs[0]
pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(pf, before=2, after=2)
make_run(pf, "React Native  •  Expo  •  TypeScript  •  Supabase PostgreSQL", size=7, color="999999")
make_run(pf, "          ", size=7, color=DARK_BROWN)
make_run(pf, "MEMENTO", size=9, color=AMBER_LIGHT, bold=True)
make_run(pf, " — Hatırla, Güvende Kal, Yaşa", size=8, color="CCCCCC")
make_run(pf, "          ", size=7, color=DARK_BROWN)
make_run(pf, "Tekirdağ NKÜ — Haziran 2026", size=7, color="999999")


# ============================================================
# SAVE
# ============================================================
output_path = os.path.expanduser("~/Desktop/Memento_Poster_A3_Pro.docx")
doc.save(output_path)
print(f"✅ Professional poster saved to: {output_path}")
